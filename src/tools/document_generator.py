"""
Document Generator Tool - Creates formatted Word documents for ERP projects
"""
from typing import Dict, List, Optional, Any
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from src.config.settings import settings
from src.utils.logger import AgentLogger

ACCENT_COLOR = RGBColor(0x1F, 0x5F, 0x4A)  # matches the frontend's pine-green accent


class DocumentGenerator:
    """Generates formatted Word (.docx) documentation for ERP projects."""

    def __init__(self):
        self.logger = AgentLogger("DocumentGenerator")
        self.output_dir = Path(settings.output_dir) / "documents"
        self.output_dir.mkdir(parents=True, exist_ok=True)

    # ------------------------------------------------------------------
    # Shared building blocks
    # ------------------------------------------------------------------

    def _new_document(self, title: str, subtitle: str) -> Document:
        doc = Document()

        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title_para.runs:
            run.font.color.rgb = ACCENT_COLOR

        subtitle_para = doc.add_paragraph(subtitle)
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_para.runs[0].italic = True
        subtitle_para.runs[0].font.size = Pt(11)

        doc.add_paragraph()
        return doc

    def _add_info_table(self, doc: Document, rows: List[tuple]):
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for label, value in rows:
            row = table.add_row().cells
            row[0].text = label
            row[0].paragraphs[0].runs[0].bold = True
            row[1].text = str(value)
        doc.add_paragraph()

    def _add_bullet_list(self, doc: Document, items: List[str]):
        if not items:
            doc.add_paragraph("None specified.", style='Intense Quote')
            return
        for item in items:
            doc.add_paragraph(str(item), style='List Bullet')

    def _add_data_table(self, doc: Document, headers: List[str], rows: List[List[str]]):
        if not rows:
            doc.add_paragraph("None specified.", style='Intense Quote')
            return
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Light Grid Accent 1'
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True
        for row_values in rows:
            row = table.add_row().cells
            for i, value in enumerate(row_values):
                row[i].text = str(value)
        doc.add_paragraph()

    def _save(self, doc: Document, prefix: str, name: str) -> str:
        safe_name = name.replace(' ', '_').replace('/', '-')
        filename = f"{prefix}_{safe_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = self.output_dir / filename
        doc.save(str(filepath))
        return str(filepath)

    # ------------------------------------------------------------------
    # Requirements document (final, from real structured requirements)
    # ------------------------------------------------------------------

    def generate_requirements_document(
        self,
        project_name: str,
        module: str,
        requirements: Dict[str, Any],
        metadata: Optional[Dict] = None
    ) -> str:
        doc = self._new_document("Requirements Document", project_name)

        self._add_info_table(doc, [
            ("Project Name", project_name),
            ("Module", module),
            ("Date", datetime.now().strftime('%Y-%m-%d')),
            ("Version", "1.0"),
            ("Status", "Draft"),
        ])

        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(requirements.get('executive_summary') or "To be completed.")

        doc.add_heading("Business Context and Objectives", level=1)
        doc.add_paragraph(requirements.get('business_context') or "To be completed.")
        doc.add_heading("Business Objectives", level=2)
        self._add_bullet_list(doc, requirements.get('objectives', []))

        doc.add_heading("Functional Requirements", level=1)
        functional_reqs = requirements.get('functional_requirements', {})
        if not functional_reqs:
            doc.add_paragraph("No functional requirements specified.", style='Intense Quote')
        for category, reqs in functional_reqs.items():
            doc.add_heading(category, level=2)
            rows = [
                [r.get('id', 'REQ-XXX'), r.get('description', ''), r.get('priority', 'Medium'),
                 r.get('acceptance_criteria', '')]
                for r in reqs
            ]
            self._add_data_table(doc, ["ID", "Description", "Priority", "Acceptance Criteria"], rows)

        doc.add_heading("Technical Requirements", level=1)
        self._add_bullet_list(doc, [r.get('description', r) if isinstance(r, dict) else r
                                     for r in requirements.get('technical_requirements', [])])

        doc.add_heading("Integration Requirements", level=1)
        self._add_bullet_list(doc, [r.get('description', r) if isinstance(r, dict) else r
                                     for r in requirements.get('integration_requirements', [])])

        doc.add_heading("Reporting Requirements", level=1)
        self._add_bullet_list(doc, [r.get('description', r) if isinstance(r, dict) else r
                                     for r in requirements.get('reporting_requirements', [])])

        doc.add_heading("Dependencies and Constraints", level=1)
        doc.add_heading("Dependencies", level=2)
        self._add_bullet_list(doc, requirements.get('dependencies', []))
        doc.add_heading("Constraints", level=2)
        self._add_bullet_list(doc, requirements.get('constraints', []))

        doc.add_heading("Assumptions", level=1)
        self._add_bullet_list(doc, requirements.get('assumptions', []))

        doc.add_heading("Approval", level=1)
        self._add_data_table(doc, ["Role", "Name", "Signature", "Date"], [
            ["Business Owner", "", "", ""],
            ["Project Manager", "", "", ""],
            ["Technical Lead", "", "", ""],
            ["Functional Consultant", "", "", ""],
        ])

        filepath = self._save(doc, "requirements", project_name)
        self.logger.log_tool_usage("generate_requirements_document", {'project': project_name, 'module': module},
                                    f"Document saved to {filepath}")
        return filepath

    # ------------------------------------------------------------------
    # Requirements questionnaire (template - no invented answers)
    # ------------------------------------------------------------------

    def generate_requirements_template(
        self,
        project_name: str,
        module: str,
        erp_system: str,
        context: Dict[str, str]
    ) -> str:
        doc = self._new_document("Requirements Gathering Questionnaire", project_name)

        self._add_info_table(doc, [
            ("Project Name", project_name),
            ("Proposed ERP System", erp_system),
            ("Module", module),
            ("Date", datetime.now().strftime('%Y-%m-%d')),
        ])

        doc.add_heading("Project Context (from intake)", level=1)
        self._add_data_table(doc, ["Field", "Value"], [
            ["Industry", context.get('industry', 'Not specified')],
            ["Organization Size", context.get('company_size', 'Not specified')],
            ["Primary Goal", context.get('primary_goal', 'Not specified')],
            ["Scope Areas", context.get('scope_areas', 'Not specified')],
        ])

        doc.add_heading("Instructions", level=1)
        doc.add_paragraph(
            "Use this questionnaire to interview stakeholders across the scope areas above. "
            "Record their actual answers - do not guess or estimate on their behalf. Once "
            "complete, bring the answers back to continue the requirements-gathering process."
        )

        sections = [
            ("1. Business Context & Objectives", [
                "What are the top 3 business drivers for this initiative?",
                "How will success be measured (KPIs, cost savings, cycle-time reductions)?",
                "What is the timeline and budget envelope for this project?",
            ]),
            ("2. Current Processes & Pain Points", [
                "Which manual or legacy processes cause the most bottlenecks today?",
                "Are there regulatory or compliance constraints that apply?",
                "What existing systems will this replace or integrate with?",
            ]),
            ("3. Functional Requirements", [
                f"For each area in scope ({context.get('scope_areas', 'the stated scope')}), "
                f"what are the must-have capabilities?",
                "What are the nice-to-have capabilities?",
                "Are there any unique workflows standard ERP modules may not cover out of the box?",
            ]),
            ("4. Data & Migration", [
                "What is the approximate volume of master and transactional data?",
                "Which data sources will need to be migrated or integrated?",
                "Are there known data quality issues in the current systems?",
            ]),
            ("5. Integration Requirements", [
                "Which external systems must this ERP connect to?",
                "What data formats or protocols are currently in use?",
            ]),
            ("6. Reporting & Analytics", [
                "What reports or dashboards are critical for each stakeholder role?",
                "Is real-time reporting required, or is periodic reporting sufficient?",
            ]),
            ("7. Constraints & Dependencies", [
                "Are there upcoming regulatory changes or infrastructure upgrades to consider?",
                "Are there vendor, staffing, or budget constraints to be aware of?",
            ]),
            ("8. Acceptance Criteria", [
                "How will each stakeholder validate that their requirements have been met?",
            ]),
        ]
        for heading, questions in sections:
            doc.add_heading(heading, level=1)
            self._add_bullet_list(doc, questions)

        doc.add_paragraph()
        note = doc.add_paragraph(
            "Once this questionnaire has been completed with real stakeholder answers, "
            "paste the responses back into the chat to continue."
        )
        note.runs[0].italic = True

        filepath = self._save(doc, "requirements_questionnaire", project_name)
        self.logger.log_tool_usage("generate_requirements_template", {'project': project_name, 'module': module},
                                    f"Template saved to {filepath}")
        return filepath

    # ------------------------------------------------------------------
    # Test cases
    # ------------------------------------------------------------------

    def generate_test_case_document(
        self,
        project_name: str,
        module: str,
        test_cases: List[Dict[str, Any]],
        test_type: str = "QA"
    ) -> str:
        doc = self._new_document(f"{test_type} Test Cases", project_name)

        self._add_info_table(doc, [
            ("Project Name", project_name),
            ("Module", module),
            ("Test Type", test_type),
            ("Date", datetime.now().strftime('%Y-%m-%d')),
            ("Total Test Cases", len(test_cases)),
        ])

        for idx, tc in enumerate(test_cases, 1):
            doc.add_heading(f"Test Case {idx}: {tc.get('scenario', 'Test Scenario')}", level=1)
            self._add_info_table(doc, [
                ("Test Case ID", tc.get('id', f'TC-{idx:03d}')),
                ("Priority", tc.get('priority', 'Medium')),
                ("Test Type", tc.get('type', 'Functional')),
            ])
            doc.add_heading("Objective", level=2)
            doc.add_paragraph(tc.get('objective', 'Not specified.'))
            doc.add_heading("Preconditions", level=2)
            self._add_bullet_list(doc, tc.get('preconditions', []))
            doc.add_heading("Test Steps", level=2)
            for step_num, step in enumerate(tc.get('steps', []), 1):
                doc.add_paragraph(f"{step_num}. {step}")
            doc.add_heading("Expected Result", level=2)
            doc.add_paragraph(tc.get('expected_result', 'Not specified.'))
            doc.add_heading("Status", level=2)
            doc.add_paragraph("☐ Pass    ☐ Fail    ☐ Blocked")

        doc.add_heading("Test Execution Summary", level=1)
        self._add_data_table(
            doc, ["Test Case ID", "Scenario", "Priority", "Status", "Tester", "Date"],
            [[tc.get('id', 'TC-XXX'), tc.get('scenario', ''), tc.get('priority', 'Medium'), "", "", ""]
             for tc in test_cases]
        )

        filepath = self._save(doc, f"test_cases_{test_type}", project_name)
        self.logger.log_tool_usage("generate_test_case_document", {'project': project_name, 'test_type': test_type},
                                    f"Document saved to {filepath}")
        return filepath

    # ------------------------------------------------------------------
    # User manual
    # ------------------------------------------------------------------

    def generate_user_manual(
        self,
        process_name: str,
        module: str,
        process_steps: List[Dict[str, Any]],
        screenshots: Optional[List[str]] = None
    ) -> str:
        doc = self._new_document(f"User Manual: {process_name}", module)

        self._add_info_table(doc, [
            ("Module", module),
            ("Process", process_name),
            ("Date", datetime.now().strftime('%Y-%m-%d')),
        ])

        doc.add_heading("Purpose", level=1)
        doc.add_paragraph(f"This manual provides step-by-step instructions for executing the "
                           f"{process_name} process in the ERP system.")

        doc.add_heading("Prerequisites", level=1)
        self._add_bullet_list(doc, [f"Access to {module} module", "Required authorizations",
                                     "Basic understanding of ERP navigation"])

        doc.add_heading("Process Steps", level=1)
        for idx, step in enumerate(process_steps, 1):
            doc.add_heading(f"Step {idx}: {step.get('title', 'Process Step')}", level=2)
            doc.add_paragraph(f"Transaction Code: {step.get('transaction', 'N/A')}").runs[0].italic = True
            doc.add_paragraph(step.get('instructions', 'Not specified.'))
            fields = step.get('fields', [])
            if fields:
                doc.add_heading("Key Fields", level=3)
                self._add_data_table(doc, ["Field", "Description", "Required", "Example"], [
                    [f.get('name', ''), f.get('description', ''), f.get('required', 'No'), f.get('example', '')]
                    for f in fields
                ])
            tips = step.get('tips')
            if tips:
                doc.add_heading("Tips", level=3)
                self._add_bullet_list(doc, tips)

        filepath = self._save(doc, "user_manual", process_name)
        self.logger.log_tool_usage("generate_user_manual", {'process': process_name, 'module': module},
                                    f"Document saved to {filepath}")
        return filepath

    # ------------------------------------------------------------------
    # Solution design
    # ------------------------------------------------------------------

    def generate_solution_design(
        self,
        project_name: str,
        module: str,
        design: Dict[str, Any]
    ) -> str:
        doc = self._new_document("Solution Design Document", project_name)

        self._add_info_table(doc, [
            ("Project Name", project_name),
            ("Module", module),
            ("Date", datetime.now().strftime('%Y-%m-%d')),
            ("Author", "ERP Consultant AI"),
        ])

        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(design.get('executive_summary') or "Not specified.")

        doc.add_heading("Solution Architecture", level=1)
        doc.add_paragraph(design.get('architecture_overview') or "Not specified.")

        doc.add_heading("Module Configuration", level=1)
        for config in design.get('configurations', []):
            doc.add_heading(config.get('component', 'Component'), level=2)
            doc.add_paragraph(config.get('description', ''))
            self._add_bullet_list(doc, config.get('steps', []))

        doc.add_heading("Integration Design", level=1)
        for integ in design.get('integrations', []):
            doc.add_heading(integ.get('name', 'Integration'), level=2)
            self._add_info_table(doc, [
                ("Type", integ.get('type', 'Real-time')),
                ("Source", integ.get('source', '')),
                ("Target", integ.get('target', '')),
            ])
            doc.add_paragraph(integ.get('description', ''))

        doc.add_heading("Customizations", level=1)
        customizations = design.get('customizations', [])
        if customizations:
            self._add_data_table(doc, ["Type", "Component", "Description", "Justification"], [
                [c.get('type', ''), c.get('component', ''), c.get('description', ''), c.get('justification', '')]
                for c in customizations
            ])
        else:
            doc.add_paragraph("No customizations required. Solution uses standard ERP functionality.")

        doc.add_heading("Migration Strategy", level=1)
        doc.add_paragraph(design.get('migration', {}).get('strategy') or "Not specified.")

        filepath = self._save(doc, "solution_design", project_name)
        self.logger.log_tool_usage("generate_solution_design", {'project': project_name, 'module': module},
                                    f"Document saved to {filepath}")
        return filepath


# Global document generator instance
doc_generator = DocumentGenerator()